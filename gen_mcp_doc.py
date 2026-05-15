import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ============ 样式设置 ============
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ============ 封面 ============
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Model Context Protocol (MCP)\n模型上下文协议')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

doc.add_paragraph('')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('技术白皮书 / 概念介绍文档')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

for _ in range(4):
    doc.add_paragraph('')

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run('2025 年 · Anthropic 发布')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ============ 目录页 ============
toc_title = doc.add_paragraph()
toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = toc_title.add_run('目  录')
run.bold = True
run.font.size = Pt(20)

doc.add_paragraph('')

toc_items = [
    ('一、什么是 MCP？', '3'),
    ('二、核心动机与解决的问题', '3'),
    ('三、MCP 架构概述', '4'),
    ('四、核心概念', '5'),
    ('五、MCP 的工作流程', '6'),
    ('六、MCP 与类似协议的对比', '7'),
    ('七、MCP 的应用场景', '8'),
    ('八、MCP 的优势', '9'),
    ('九、如何开始使用 MCP', '9'),
    ('十、总结与展望', '10'),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Cm(14.5))
    run = p.add_run(f'{item}')
    run.font.size = Pt(12)

doc.add_page_break()

# ============ 正文 ============

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(22)
    return p

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(20)
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # Add shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    p._element.get_or_add_pPr().append(shading)
    return p

# ============ 第一章 ============
add_heading('一、什么是 MCP？', 1)

add_body('Model Context Protocol（MCP，模型上下文协议）是由 Anthropic 提出并开源的一种开放标准协议，旨在为大型语言模型（LLM）提供一种统一、安全且高效的方式来连接和访问外部数据源、工具与服务。')

add_body('可以将 MCP 理解为 AI 应用程序的 "USB-C 接口" —— 它定义了一套标准的通信规范，使得 AI 模型能够通过统一的接口与各种外部系统（数据库、文件系统、API、Web 服务等）进行交互，而无需为每个集成单独开发适配方案。')

add_body('MCP 于 2024 年底由 Anthropic 首次公布，并迅速获得了包括 OpenAI、Google、微软、JetBrains 等在内的众多行业巨头的支持，成为 AI Agent 领域事实上的标准协议之一。')

# ============ 第二章 ============
add_heading('二、核心动机与解决的问题', 1)

add_heading('2.1 当前 AI 集成的痛点', 2)

add_body('在 MCP 出现之前，将 LLM 连接到外部数据源通常面临以下挑战：')

add_bullet('碎片化集成：每个数据源（数据库、API、文件系统）需要单独的适配代码，导致大量重复工作。')
add_bullet('安全性难以保障：每个集成各自实现认证、授权机制，容易产生安全漏洞。')
add_bullet('上下文受限：LLM 只能依赖训练数据或手动提供的上下文，无法实时访问最新或私有的数据。')
add_bullet('工具定义不统一：不同模型和框架对工具（Tool）的定义和调用方式各不相同，缺乏互操作性。')

add_heading('2.2 MCP 的解决思路', 2)

add_body('MCP 通过以下方式解决上述问题：')

add_bullet('标准化接口：定义统一的协议规范，任何兼容 MCP 的客户端和服务端都可以无缝通信。')
add_bullet('关注点分离：将 AI 模型与外部系统的连接逻辑抽象为独立的服务，降低耦合。')
add_bullet('安全内建：协议层内置了权限控制、资源隔离等安全机制。')
add_bullet('即插即用：开发者可以像安装 USB 设备一样，通过 MCP Server 快速为 AI 应用添加新能力。')

# ============ 第三章 ============
add_heading('三、MCP 架构概述', 1)

add_body('MCP 采用经典的 客户端-服务器（Client-Server）架构，包含三个核心角色：')

add_heading('3.1 架构组件', 2)

# 表格
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['组件', '角色', '说明']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True

data = [
    ['MCP Host', '宿主应用', '运行 MCP 的 AI 应用，如 Claude Desktop、IDE 插件等'],
    ['MCP Client', '客户端', '与 MCP Server 建立一对一连接的通信端点'],
    ['MCP Server', '服务端', '提供特定能力（工具、资源、提示模板）的轻量级服务'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, cell_data in enumerate(row_data):
        table.rows[row_idx].cells[col_idx].text = cell_data

doc.add_paragraph('')

add_heading('3.2 架构示意图', 2)

add_code_block('┌─────────────────────────────────────────┐')
add_code_block('│              MCP Host                  │')
add_code_block('│  (Claude Desktop / IDE / App)          │')
add_code_block('│  ┌─────────────┐  ┌─────────────┐     │')
add_code_block('│  │ MCP Client 1 │  │ MCP Client 2 │     │')
add_code_block('│  └──────┬──────┘  └──────┬──────┘     │')
add_code_block('└─────────┼────────────────┼────────────┘')
add_code_block('          │                │               ')
add_code_block('          ▼                ▼               ')
add_code_block('┌─────────────────┐  ┌─────────────────┐  ')
add_code_block('│   MCP Server 1  │  │   MCP Server 2  │  ')
add_code_block('│ (数据库连接器)   │  │ (文件系统连接器)  │  ')
add_code_block('└─────────────────┘  └─────────────────┘  ')

add_heading('3.3 传输层', 2)

add_body('MCP 支持两种传输方式：')

add_bullet('stdio 传输：通过标准输入/输出进行通信，适用于本地进程间通信。')
add_bullet('SSE（Server-Sent Events）传输：通过 HTTP 进行通信，适用于远程服务。')

# ============ 第四章 ============
add_heading('四、核心概念', 1)

add_heading('4.1 资源（Resources）', 2)
add_body('资源是 MCP Server 暴露给 AI 模型的数据源。资源可以是文件、数据库记录、API 响应等。MCP 使用 URI 方案来标识和定位资源，例如：')
add_code_block('file:///home/user/data/report.pdf')
add_code_block('db://users/123/profile')
add_code_block('api://weather/beijing/current')

add_heading('4.2 工具（Tools）', 2)
add_body('工具是 MCP Server 提供的可调用功能，类似于 API 端点。模型可以动态发现并调用这些工具来执行操作。工具通常使用 JSON Schema 来描述其输入参数。')

add_heading('4.3 提示模板（Prompts）', 2)
add_body('提示模板是预定义的提示词模板，允许 Server 为特定的交互场景提供结构化的提示信息。这有助于模型理解如何与 Server 的能力进行交互。')

add_heading('4.4 采样（Sampling）', 2)
add_body('采样允许 Server 请求 Host（宿主应用）调用 LLM 生成文本。这是一种反向通信机制，使得 Server 能够利用模型的生成能力来满足用户的需求。')

# ============ 第五章 ============
add_heading('五、MCP 的工作流程', 1)

add_heading('5.1 连接建立', 2)
add_body('MCP Client 与 MCP Server 建立连接的过程分为两个阶段：')

add_body('阶段一 — 初始化：')
add_bullet('Client 向 Server 发送初始化请求，包含协议版本和客户端能力声明。')
add_bullet('Server 回复其支持的协议版本和服务端能力声明。')

add_body('阶段二 — 能力协商：')
add_bullet('双方交换能力信息，确定可用的资源、工具和提示模板。')
add_bullet('完成握手后进入正常通信阶段。')

add_heading('5.2 消息交互', 2)
add_body('MCP 使用 JSON-RPC 2.0 作为消息协议，支持三种消息类型：')

add_bullet('请求（Request）：Client 向 Server 发起的调用，期望获得响应。')
add_bullet('通知（Notification）：单向消息，不需要响应。')
add_bullet('响应（Response）：对请求的回复。')

add_heading('5.3 典型交互流程', 2)
add_code_block('1. 用户向 AI 应用提问："帮我查询上个月的销售数据"')
add_code_block('2. 模型分析意图，决定调用数据库 MCP Server')
add_code_block('3. MCP Client 向 MCP Server 发送查询请求')
add_code_block('4. MCP Server 执行数据库查询，返回结果')
add_code_block('5. 模型基于返回的数据生成回答')
add_code_block('6. 应用将回答呈现给用户')

# ============ 第六章 ============
add_heading('六、MCP 与类似协议的对比', 1)

# 对比表
table2 = doc.add_table(rows=5, cols=4)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ['特性', 'MCP', 'Function Calling', 'Plugin / Tool']
for i, header in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = header
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True

comparison = [
    ['标准化程度', '开放标准协议', '各模型各自实现', '平台绑定'],
    ['传输协议', 'JSON-RPC 2.0', '无标准协议', 'HTTP / 自定义'],
    ['动态发现', '支持（资源/工具发现）', '需预定义', '部分支持'],
    ['安全模型', '内建权限控制', '需自行实现', '平台提供'],
]
for row_idx, row_data in enumerate(comparison, 1):
    for col_idx, cell_data in enumerate(row_data):
        table2.rows[row_idx].cells[col_idx].text = cell_data

doc.add_paragraph('')

# ============ 第七章 ============
add_heading('七、MCP 的应用场景', 1)

add_heading('7.1 智能编程助手', 2)
add_body('IDE 插件可以通过 MCP 连接代码库、文档、API 参考和项目管理工具，为开发者提供上下文感知的编程辅助。例如：自动分析代码库结构、实时获取 API 文档、管理 Git 操作等。')

add_heading('7.2 数据分析与 BI', 2)
add_body('AI 数据助手可以通过 MCP 连接数据库、数据仓库和 BI 工具，实现自然语言查询数据、自动生成报表和可视化图表。')

add_heading('7.3 自动化工作流', 2)
add_body('通过 MCP 连接各种 SaaS 工具（Slack、Notion、Jira、GitHub 等），实现跨应用的自动化工作流，例如自动创建工单、更新文档、发送通知等。')

add_heading('7.4 企业知识库', 2)
add_body('MCP 可以连接企业内部的知识管理系统、文档库和 Wiki，使 AI 能够基于企业私有知识回答问题，成为智能的企业知识助手。')

add_heading('7.5 个人助理', 2)
add_body('连接日历、邮件、通讯录、笔记等个人应用，实现智能日程管理、邮件起草、信息整理等个人事务的自动化处理。')

# ============ 第八章 ============
add_heading('八、MCP 的优势', 1)

add_bullet('开放标准：由 Anthropic 发起并开源，社区驱动的标准协议，不绑定特定平台或模型。')
add_bullet('广泛生态支持：获得 OpenAI、Google、Microsoft、JetBrains 等行业巨头的支持。')
add_bullet('即插即用：丰富的社区 MCP Server 生态，可快速为 AI 应用添加新能力。')
add_bullet('安全可控：协议层内建权限控制和资源隔离机制。')
add_bullet('双向通信：支持 Client→Server 和 Server→Host 的双向通信模式。')
add_bullet('动态发现：MCP Server 可以动态声明其能力，Client 无需预先配置。')

# ============ 第九章 ============
add_heading('九、如何开始使用 MCP', 1)

add_heading('9.1 官方资源', 2)
add_bullet('GitHub 仓库：github.com/modelcontextprotocol')
add_bullet('官方文档：modelcontextprotocol.io')
add_bullet('协议规范：spec.modelcontextprotocol.io')

add_heading('9.2 SDK 支持', 2)
add_body('Anthropic 官方提供以下语言的 SDK：')
add_bullet('Python SDK：用于构建 MCP Client 和 Server')
add_bullet('TypeScript SDK：用于 Node.js 环境的 MCP 开发')
add_bullet('Java SDK：用于 Java 生态的 MCP 开发')
add_bullet('Kotlin SDK：用于 Kotlin 生态的 MCP 开发')

add_heading('9.3 快速开始示例', 2)
add_body('一个简单的 MCP Server 示例（Python）：')

add_code_block('from mcp.server import Server, NotificationOptions')
add_code_block('from mcp.server.models import InitializationOptions')
add_code_block('')
add_code_block('# 创建 MCP Server')
add_code_block('server = Server("example-server")')
add_code_block('')
add_code_block('# 定义一个工具')
add_code_block('@server.list_tools()')
add_code_block('async def handle_list_tools() -> list[Tool]:')
add_code_block('    return [')
add_code_block('        Tool(')
add_code_block('            name="greet",')
add_code_block('            description="向用户问好",')
add_code_block('            inputSchema={"type": "object",')
add_code_block('                          "properties": {"name": {"type": "string"}}}')
add_code_block('        )')
add_code_block('    ]')

# ============ 第十章 ============
add_heading('十、总结与展望', 1)

add_body('Model Context Protocol (MCP) 代表了 AI 集成领域的一个重要里程碑。通过提供统一、开放的标准协议，MCP 极大地简化了 AI 模型与外部系统的连接过程，为 AI Agent 的发展奠定了坚实的基础设施。')

add_body('随着 AI 技术从"对话式 AI"向"自主 Agent"的演进，MCP 作为连接 AI 模型与外部世界的"桥梁"，其重要性将日益凸显。可以预见，MCP 将成为 AI 开发生态中的核心组件，类似于 HTTP 在 Web 世界中的地位。')

add_body('未来，MCP 生态将持续扩展，更多的 MCP Server 将被开发出来，覆盖更广泛的应用场景。同时，MCP 协议本身也将不断演进，支持更丰富的通信模式、更高效的传输机制和更完善的安全模型。')

# 分隔线
doc.add_paragraph('')
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 文档结束 —')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.size = Pt(11)

# ============ 保存文档 ============
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'MCP_Model_Context_Protocol_介绍.docx')
doc.save(output_path)
print(f'文档已生成: {output_path}')
